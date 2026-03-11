"""
공사명칭 검토 엔진 모듈.

목적:
    보고서 파일(PDF/HWP/XLSX 등)에서 추출한 텍스트를
    마스터 DB의 공식 공사명칭과 비교하여 일치/경고/불일치를 판정한다.

주요 기능:
    - 파일 형식별 텍스트 추출
    - 마스터 명칭 로드(내장/서버)
    - 접두어 분리 기반 보정 매칭
    - 검토 결과 Excel 리포트 생성

불일치 사유 3분류:
    - 공사명 불일치: 순수명칭 자체가 다름
    - 접두어(사업유형) 불일치: 순수명칭은 맞으나 접두어가 다름
    - 연도 불일치: 접두어·공사명은 같고 연도만 다름

버전:
    v3.2
"""

import os
import re
import bisect
import queue
import threading
from multiprocessing import Pool, freeze_support
from datetime import datetime
from difflib import SequenceMatcher
from typing import Optional

# PDF 병렬 추출 워커 수 (CPU 코어 기반, 최대 8)
_CPU_COUNT = os.cpu_count() or 4
PDF_POOL_WORKERS = max(4, min(_CPU_COUNT // 2, 8))

# 마지막 마스터 로드 실패 사유(서버 동기화 실패 메시지 보관)
_LAST_MASTER_LOAD_ERROR = ""

# filepath -> {page_num -> {line_num -> (x0, top, bottom)}}
_PDF_FIRST_CHAR_COORDS: dict[str, dict[int, dict[int, tuple[float, float, float]]]] = {}

# Google Sheets CSV export 등에서 내려오는 헤더 행 후보
MASTER_NAME_HEADERS = {"공사약식명", "공사명", "공사명칭"}


# ═══════════════════════════════════════════════════════════
#  파일 텍스트 추출 (지연 import)
# ═══════════════════════════════════════════════════════════
# 지원하는 입력 파일 확장자 목록
SUPPORTED_EXTENSIONS = {'.xlsx', '.pdf', '.docx', '.hwp', '.hwpx', '.csv'}

# 마스터DB 공식명칭에 사용되는 접두어 목록
# 긴 접두어부터 매칭하기 위해 사용 시 길이 역순 정렬이 필요하다.
KNOWN_PREFIXES: tuple[str, ...] = (
    "민참.분양", "자체.시공", "자체.시행", "자체.공모", "시행도급",
    "순수내역", "민간", "종평", "적격", "기본", "종심", "턴키",
    "해외", "통합", "도로", "영업", "도전", "도운", "국운",
    "지운", "민운", "민임", "실시", "자체", "민참",
    "CM", "BOT", "BTL", "BTO", "ITS",
)

# PDF/HWP 추출 시 명칭 앞에 붙는 특수기호
STRIP_CHARS = "▶□■❑◆●○◇△▽☆★※→←↑↓·"

# 공식명칭 접두어 파싱용 정규식
PREFIX_PATTERN = re.compile(r'^\(([^)]*)\)\s*(.*)$')

# 명칭 앞뒤에 붙어도 경고로 보지 않을 단순 기호/조사 패턴
_STRIP_SURROUNDING = re.compile(
    r'^[『「\["\s]+'
    r'|'
    r'[』」\],\."外(（\s]+$'
    r'|'
    r'(?:은|는|이|가|을|를|의|과|와|도|만)+$'
)

# 순수명칭 내 연도 숫자 패턴 (예: "26년", "2026년")
_YEAR_NUM_RE = re.compile(r"(\d{2,4})\s*년")


def _is_year_only_diff(bare_a: str, bare_b: str) -> bool:
    """두 순수명칭이 연도 부분만 다른지 판별한다.

    Args:
        bare_a: 첫 번째 순수명칭.
        bare_b: 두 번째 순수명칭.

    Returns:
        연도 숫자만 다르고 나머지가 동일하면 True.
    """
    placeholder = "\x00YEAR\x00"
    norm_a = _YEAR_NUM_RE.sub(placeholder, bare_a)
    norm_b = _YEAR_NUM_RE.sub(placeholder, bare_b)
    if norm_a != norm_b:
        return False
    years_a = _YEAR_NUM_RE.findall(bare_a)
    years_b = _YEAR_NUM_RE.findall(bare_b)
    return years_a != years_b


def _strip_surrounding(name: str) -> str:
    """명칭 바깥의 단순 기호/조사를 제거한다."""
    cleaned = str(name or "").strip()
    previous = None
    while cleaned != previous:
        previous = cleaned
        cleaned = _STRIP_SURROUNDING.sub("", cleaned).strip()
    return cleaned


def _get_attached_side_segment(text: str, index: int, is_left: bool) -> str:
    """명칭 좌우에 공백 없이 붙은 연속 토큰을 추출한다."""
    source = text[:index] if is_left else text[index:]
    if not source:
        return ""

    if is_left:
        match = re.search(r"\S+$", source)
    else:
        match = re.match(r"^\S+", source)
    return match.group(0) if match else ""


def extract_from_xlsx(filepath: str) -> list:
    """엑셀 파일에서 셀 텍스트를 추출한다.

    Args:
        filepath: `.xlsx` 파일 경로.

    Returns:
        `(위치, 텍스트)` 튜플 목록.
    """
    from openpyxl import load_workbook
    texts = []
    wb = load_workbook(filepath, data_only=True)
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=False):
            for cell in row:
                if cell.value is not None:
                    val = str(cell.value).strip()
                    if val:
                        loc = f"[{ws.title}] {cell.coordinate}"
                        texts.append((loc, val))
    wb.close()
    return texts


def _extract_single_page_pdfplumber(args):
    """단일 페이지를 pdfplumber로 추출한다.

    Args:
        args: `(filepath, page_num)` 튜플. `page_num`은 1-based.

    Returns:
        `(page_num, [(위치, 텍스트), ...], {line_num: (x0, top, bottom)})` 튜플.
        좌표는 순수명칭(접두어 제거 후)의 첫 한글/영문/숫자 글자 기준이다.
    """
    filepath, page_num = args
    import pdfplumber

    _search_char_re = re.compile(r'[가-힣a-zA-Z0-9]')

    results = []
    first_char_coords: dict[int, tuple[float, float, float]] = {}
    try:
        with pdfplumber.open(filepath) as pdf:
            page_idx = page_num - 1
            if 0 <= page_idx < len(pdf.pages):
                page = pdf.pages[page_idx]
                page_text = page.extract_text() or ""

                lines = page_text.splitlines()
                for line_num, line in enumerate(lines, 1):
                    stripped = line.strip()
                    if stripped:
                        results.append((f"P{page_num} L{line_num}", stripped))

                chars = page.chars
                if not chars:
                    return page_num, results, first_char_coords

                char_idx = 0
                for line_num, line in enumerate(lines, 1):
                    stripped = line.strip()
                    if not stripped:
                        continue

                    # 접두어를 제거하여 순수명칭의 시작 위치를 파악한다.
                    _, bare = strip_known_prefix(stripped)
                    if not bare:
                        bare = stripped

                    # 순수명칭에서 첫 한글/영문/숫자 글자를 추출한다.
                    m = _search_char_re.search(bare)
                    if not m:
                        continue
                    target_char = m.group(0)

                    # bare 앞 3글자(한글/영문/숫자만)를 연속 매칭 검증용으로 수집한다.
                    verify_chars = []
                    for ch in bare:
                        if _search_char_re.match(ch):
                            verify_chars.append(ch)
                            if len(verify_chars) >= 3:
                                break

                    # chars 배열에서 target_char를 찾되, 후속 글자 연속 매칭으로 검증한다.
                    best_i = None
                    best_score = 0
                    for i in range(char_idx, len(chars)):
                        if chars[i]["text"] != target_char:
                            continue

                        score = 1
                        for offset, expected in enumerate(verify_chars[1:], 1):
                            ni = i + offset
                            if ni < len(chars) and chars[ni]["text"] == expected:
                                score += 1
                            else:
                                break

                        if score > best_score:
                            best_score = score
                            best_i = i

                        # 3글자 모두 매칭되면 확정
                        if best_score >= len(verify_chars):
                            break

                    if best_i is not None:
                        c = chars[best_i]
                        first_char_coords[line_num] = (
                            float(c["x0"]),
                            float(c["top"]),
                            float(c["bottom"]),
                        )
                        char_idx = best_i + 1

    except Exception:
        pass
    return page_num, results, first_char_coords


def _reconstruct_prefix(bare: str, lines: list, line_idx: int,
                        master_set: set, known_prefixes: tuple) -> str | None:
    """PyMuPDF 텍스트 분리로 인해 접두어가 누락된 순수명칭을 재조합한다.

    Args:
        bare: 순수명칭 문자열.
        lines: 해당 페이지의 전체 줄 목록.
        line_idx: bare가 위치한 줄 인덱스.
        master_set: 공식명칭 집합.
        known_prefixes: 알려진 접두어 튜플.

    Returns:
        재조합된 공식명칭 또는 None.
    """
    cur = lines[line_idx].strip()
    sorted_pfx = sorted(known_prefixes, key=len, reverse=True)

    # 1) 위/아래 3줄까지 (접두어) 형태 탐색
    for dist in range(1, 4):
        for idx in [line_idx + dist, line_idx - dist]:
            if 0 <= idx < len(lines):
                ln = lines[idx].strip()
                paren = re.search(r'\(([^)]+)\)', ln)
                if paren:
                    raw = paren.group(1)
                    for variant in [raw, raw.replace(',', '.')]:
                        candidate = '(' + variant + ')' + bare
                        if candidate in master_set:
                            return candidate

    # 2) 접두어가 괄호 없이 순수명칭 앞에 직접 붙은 형태
    for pfx in sorted_pfx:
        if cur.startswith(pfx + bare):
            candidate = '(' + pfx + ')' + bare
            if candidate in master_set:
                return candidate

    # 3) ')순수명칭' 형태 (접두어가 윗줄에 분리)
    if cur.startswith(')') and bare in cur:
        for back in range(1, min(5, line_idx + 1)):
            combined = ''.join(
                lines[line_idx - j].strip() for j in range(back, 0, -1)
            )
            for pfx in sorted_pfx:
                if pfx in combined or pfx.replace('.', ',') in combined:
                    candidate = '(' + pfx + ')' + bare
                    if candidate in master_set:
                        return candidate

    # 4) 현재 줄에서 (접두어)순수명칭 패턴 직접 탐색
    paren_match = re.search(r'\(([^)]+)\)\s*' + re.escape(bare), cur)
    if paren_match:
        raw = paren_match.group(1)
        for variant in [raw, raw.replace(',', '.')]:
            candidate = '(' + variant + ')' + bare
            if candidate in master_set:
                return candidate

    return None


def extract_from_pdf(filepath: str, progress_callback=None) -> list:
    """PDF 파일에서 줄 단위 텍스트를 추출한다.

    Args:
        filepath: .pdf 파일 경로.
        progress_callback: PDF 추출 단계 진행률 콜백(0.0~0.3 구간).

    Returns:
        (위치, 텍스트) 튜플 목록.
    """
    def _emit_progress(value: float):
        """PDF 추출 진행률 콜백을 안전하게 호출한다."""
        if progress_callback is None:
            return
        try:
            clamped = max(0.0, min(0.3, float(value)))
            progress_callback(clamped)
        except Exception:
            pass

    texts: list[tuple[str, str]] = []
    _PDF_FIRST_CHAR_COORDS[filepath] = {}

    try:
        import pdfplumber

        with pdfplumber.open(filepath) as pdf:
            total_pages = len(pdf.pages)
            if total_pages == 0:
                return []

        _emit_progress(0.05)

        try:
            page_args = [(filepath, pg_num) for pg_num in range(1, total_pages + 1)]
            with Pool(PDF_POOL_WORKERS) as pool:
                page_results = pool.map(_extract_single_page_pdfplumber, page_args)

            page_results.sort(key=lambda item: item[0])

            parallel_texts: list[tuple[str, str]] = []
            coords_by_page: dict[int, dict[int, tuple[float, float, float]]] = {}
            for page_num, lines, first_char_coords in page_results:
                parallel_texts.extend(lines)
                if first_char_coords:
                    coords_by_page[page_num] = first_char_coords
                _emit_progress(0.05 + 0.25 * (page_num / total_pages))

            _emit_progress(0.3)
            if parallel_texts:
                _PDF_FIRST_CHAR_COORDS[filepath] = coords_by_page
                return parallel_texts
        except Exception:
            pass
    except Exception:
        pass

    try:
        import fitz
    except ImportError:
        raise ImportError("pdfplumber/PyMuPDF 둘 다 사용할 수 없습니다.")

    doc = fitz.open(filepath)
    try:
        total_pages = len(doc)
        if total_pages == 0:
            return []

        _emit_progress(0.05)
        for pg_idx in range(total_pages):
            page = doc[pg_idx]
            page_text = page.get_text()
            page_num = pg_idx + 1

            if page_text:
                for line_num, line in enumerate(page_text.splitlines(), 1):
                    stripped = line.strip()
                    if stripped:
                        location = f"P{page_num} L{line_num}"
                        texts.append((location, stripped))

            _emit_progress(0.05 + 0.25 * ((pg_idx + 1) / total_pages))
    finally:
        doc.close()

    _emit_progress(0.3)
    return texts


def extract_from_docx(filepath: str) -> list:
    """DOCX 파일에서 문단/표 텍스트를 추출한다.

    Args:
        filepath: `.docx` 파일 경로.

    Returns:
        `(위치, 텍스트)` 튜플 목록.
    """
    from docx import Document
    texts = []
    doc = Document(filepath)
    for i, para in enumerate(doc.paragraphs, 1):
        t = para.text.strip()
        if t:
            texts.append((f"문단{i}", t))
    for ti, table in enumerate(doc.tables, 1):
        for ri, row in enumerate(table.rows, 1):
            for ci, cell in enumerate(row.cells, 1):
                t = cell.text.strip()
                if t:
                    texts.append((f"표{ti}({ri},{ci})", t))
    return texts


def extract_from_hwp(filepath: str) -> list:
    """HWP(OLE) 파일에서 본문 텍스트를 추출한다.

    Args:
        filepath: `.hwp` 파일 경로.

    Returns:
        `(위치, 텍스트)` 튜플 목록.

    Raises:
        RuntimeError: 파일 파싱 또는 본문 추출에 실패한 경우.
    """
    import struct
    import zlib

    try:
        import olefile
    except ImportError as e:
        raise ImportError("HWP 파싱을 위해 olefile 패키지가 필요합니다.") from e

    texts = []
    line_num = 0

    try:
        if not olefile.isOleFile(filepath):
            raise ValueError("OLE 형식의 .hwp 파일이 아닙니다.")

        with olefile.OleFileIO(filepath) as hwp:
            if not hwp.exists("FileHeader"):
                raise ValueError("FileHeader 스트림을 찾지 못했습니다.")

            header_data = hwp.openstream("FileHeader").read()
            is_compressed = (
                len(header_data) > 36 and (header_data[36] & 1) == 1
            )

            section_nums = []
            for entry in hwp.listdir(streams=True, storages=False):
                if (
                    len(entry) >= 2
                    and entry[0] == "BodyText"
                    and entry[1].startswith("Section")
                ):
                    try:
                        section_nums.append(
                            int(entry[1].replace("Section", ""))
                        )
                    except ValueError:
                        continue
            section_nums = sorted(set(section_nums))

            for snum in section_nums:
                stream_name = f"BodyText/Section{snum}"
                if not hwp.exists(stream_name):
                    continue

                data = hwp.openstream(stream_name).read()
                if is_compressed:
                    try:
                        data = zlib.decompress(data, -15)
                    except zlib.error:
                        continue

                i = 0
                data_len = len(data)
                while i + 4 <= data_len:
                    header_val = struct.unpack_from("<I", data, i)[0]
                    rec_type = header_val & 0x3FF
                    rec_len = (header_val >> 20) & 0xFFF
                    i += 4

                    if rec_len == 0xFFF:
                        if i + 4 > data_len:
                            break
                        rec_len = struct.unpack_from("<I", data, i)[0]
                        i += 4

                    if rec_len < 0 or i + rec_len > data_len:
                        break

                    if rec_type == 67:
                        try:
                            raw = data[i:i + rec_len]
                            text = raw.decode("utf-16-le", errors="ignore")

                            clean_chars = []
                            for ch in text:
                                if ch in ("\r", "\n"):
                                    continue
                                if ord(ch) < 32:
                                    continue
                                clean_chars.append(ch)

                            clean = "".join(clean_chars).strip()
                            if clean:
                                line_num += 1
                                texts.append((f"HWP L{line_num}", clean))
                        except Exception:
                            pass

                    i += rec_len
    except Exception as e:
        raise RuntimeError(f"HWP 읽기 실패: {e}") from e

    if not texts:
        raise RuntimeError("HWP 읽기 실패: 본문 텍스트를 찾지 못했습니다.")
    return texts


def extract_from_hwpx(filepath: str) -> list:
    """HWPX(ZIP+XML) 파일에서 본문 텍스트를 추출한다.

    Args:
        filepath: `.hwpx` 파일 경로.

    Returns:
        `(위치, 텍스트)` 튜플 목록.

    Raises:
        RuntimeError: 파일 열기 또는 XML 파싱에 실패한 경우.
    """
    import zipfile
    from xml.etree import ElementTree as ET

    texts = []
    line_num = 0

    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            section_files = sorted([
                name for name in zf.namelist()
                if name.startswith('Contents/section') and name.endswith('.xml')
            ])

            if not section_files:
                raise ValueError("Contents/section*.xml을 찾지 못했습니다.")

            for section_file in section_files:
                with zf.open(section_file) as sf:
                    root = ET.parse(sf).getroot()

                    for elem in root.iter():
                        tag = elem.tag
                        if isinstance(tag, str) and '}' in tag:
                            tag = tag.split('}', 1)[1]
                        if tag != 't' or not elem.text:
                            continue

                        text = elem.text.strip()
                        if not text:
                            continue
                        line_num += 1
                        texts.append((f"HWPX L{line_num}", text))

        if not texts:
            raise ValueError("본문 텍스트를 찾지 못했습니다.")
        return texts
    except Exception as e:
        raise RuntimeError(f"HWPX 파일 읽기 실패: {e}")


def extract_from_csv(filepath: str) -> list:
    """CSV 파일을 다중 인코딩으로 읽어 텍스트를 추출한다.

    Args:
        filepath: `.csv` 파일 경로.

    Returns:
        `(위치, 텍스트)` 튜플 목록.

    Raises:
        ValueError: 지원 인코딩으로 디코딩할 수 없는 경우.
    """
    import csv

    texts = []
    for encoding in ['utf-8', 'cp949', 'euc-kr']:
        try:
            with open(filepath, 'r', encoding=encoding, newline='') as f:
                reader = csv.reader(f)
                for ri, row in enumerate(reader, 1):
                    for ci, val in enumerate(row, 1):
                        val = val.strip()
                        if val:
                            texts.append((f"R{ri} C{ci}", val))
            return texts
        except UnicodeDecodeError:
            continue

    raise ValueError(f"CSV 인코딩 인식 불가: {filepath}")


def extract_text_from_file(
    filepath: str,
    include_pdf_tables: bool = False,
    progress_callback=None,
) -> list:
    """파일 확장자에 따라 적절한 텍스트 추출기를 호출한다.

    Args:
        filepath: 입력 파일 경로.
        include_pdf_tables: 호환용 인자(현재 미사용).
        progress_callback: 텍스트 추출 단계 진행률 콜백.

    Returns:
        `(위치, 텍스트)` 튜플 목록.
    """
    _ = include_pdf_tables
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_from_pdf(filepath, progress_callback=progress_callback)

    dispatch = {
        ".xlsx": extract_from_xlsx,
        ".docx": extract_from_docx,
        ".hwp": extract_from_hwp,
        ".hwpx": extract_from_hwpx,
        ".csv": extract_from_csv,
    }
    func = dispatch.get(ext)
    if func is None:
        raise ValueError(f"지원하지 않는 형식: {ext}")
    return func(filepath)


def fetch_master_from_server(url: str, timeout: int = 10) -> tuple[list, Optional[str]]:
    """서버 CSV를 내려받아 마스터 명칭 목록을 로드한다.

    Args:
        url: 마스터 CSV URL.
        timeout: HTTP 타임아웃(초).

    Returns:
        `(정제된 명칭 목록, DB 기준일)` 튜플.
    """
    import csv
    import io
    from email.utils import parsedate_to_datetime
    import requests

    response = requests.get(url, timeout=timeout)
    response.raise_for_status()

    db_date = None
    last_modified = response.headers.get("Last-Modified")
    if last_modified:
        try:
            db_date = parsedate_to_datetime(last_modified).strftime("%Y-%m-%d")
        except (TypeError, ValueError, OverflowError):
            db_date = None

    raw_text = _decode_master_csv_text(response.content)
    reader = csv.reader(io.StringIO(raw_text))
    raw_names = [row[0] for row in reader if row]
    names = _sanitize_master_names(raw_names)
    if len(names) < 10:
        raise ValueError("서버 데이터가 너무 적습니다")

    return names, db_date


# ═══════════════════════════════════════════════════════════
#  기본 마스터 명칭
# ═══════════════════════════════════════════════════════════
def _sanitize_master_names(raw_names: list) -> list:
    """마스터 명칭 원본 목록을 중복 제거/정제한다.

    Args:
        raw_names: 원본 명칭 목록.

    Returns:
        공백/중복/`None`이 제거된 명칭 목록.
    """
    sanitized = []
    seen = set()
    for raw in raw_names:
        if raw is None:
            continue
        name = str(raw).strip().lstrip("\ufeff")
        if name in MASTER_NAME_HEADERS:
            continue
        if not name or name in seen:
            continue
        seen.add(name)
        sanitized.append(name)
    return sanitized


def _decode_master_csv_text(raw: bytes) -> str:
    """마스터 CSV 바이트를 가능한 인코딩으로 디코딩한다.

    Args:
        raw: CSV 원본 바이트.

    Returns:
        디코딩된 문자열.
    """
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _normalize_known_prefix_punctuation(text: str) -> str:
    """선두 접두어 괄호 안 쉼표를 마침표로 정규화한다.

    Args:
        text: 원본 후보 문자열.

    Returns:
        KNOWN_PREFIXES에 해당하는 접두어만 쉼표/마침표 표기를 통일한 문자열.
    """
    clean = str(text or "").strip()
    pm = PREFIX_PATTERN.match(clean)
    if not pm:
        return clean

    prefix = pm.group(1).strip()
    normalized_prefix = re.sub(r"\s*,\s*", ".", prefix)
    if normalized_prefix == prefix or normalized_prefix not in KNOWN_PREFIXES:
        return clean

    bare = pm.group(2).strip()
    return f"({normalized_prefix}){bare}"


def strip_known_prefix(text: str) -> tuple[str, str]:
    """텍스트 앞에 붙은 알려진 접두어를 분리한다.

    Args:
        text: 검사 대상 문자열.

    Returns:
        `(접두어, 순수명칭)` 튜플. 접두어 미발견 시 `("", 원본)` 반환.
    """
    raw = str(text or "")
    cleaned = raw.strip().strip(STRIP_CHARS).strip()
    cleaned = _normalize_known_prefix_punctuation(cleaned)
    if not cleaned:
        return "", ""

    pm = PREFIX_PATTERN.match(cleaned)
    if pm:
        prefix = pm.group(1).strip()
        bare = pm.group(2).strip()
        if prefix and bare:
            return prefix, bare
        return "", cleaned

    for prefix in sorted(KNOWN_PREFIXES, key=len, reverse=True):
        if cleaned.startswith(prefix):
            bare = cleaned[len(prefix):].strip()
            if bare:
                return prefix, bare

    return "", cleaned


def split_official_name(name: str) -> tuple[str, str]:
    """공식명칭을 `(접두어, 순수명칭)`으로 분리한다.

    Args:
        name: 마스터 공식명칭.

    Returns:
        `(접두어, 순수명칭)` 튜플.
    """
    clean = _normalize_known_prefix_punctuation(str(name or "").strip())
    pm = PREFIX_PATTERN.match(clean)
    if pm:
        return pm.group(1).strip(), pm.group(2).strip()
    return strip_known_prefix(clean)


def build_bare_name_index(master_names: list) -> dict[str, list[str]]:
    """마스터 명칭으로 순수명칭 인덱스를 생성한다.

    Args:
        master_names: 공식명칭 목록.

    Returns:
        `순수명칭 -> [공식명칭, ...]` 딕셔너리.
    """
    index: dict[str, list[str]] = {}
    for official in master_names:
        _, bare = split_official_name(official)
        if not bare:
            continue
        index.setdefault(bare, []).append(official)
    return index


DEFAULT_MASTER_NAMES = [
    "(민간)부산미음동물류",
    "(CM)평택고덕10A",
    "(종평)안심뉴타운A",
    "(적격)서해선103역사",
    "(민간)안성CDC물류",
    "(민간)제주한화우주센터",
    "(민간)쿠팡김천물류",
    "(민간)전주대자인병원",
    "(BOT)고삼호수 휴게소",
    "(기본)이천호국원확충",
    "(종심)운정역환승주차장",
    "(BTL)춘천기계공고",
    "(BTL)춘천기계공고(봄내중)",
    "(적격)고양주차장전기",
    "(종심)의왕월암1A",
    "(종심)경산대임1A",
    "(종심)인천계양5A",
    "(종심)성남복정1A",
    "(적격)강원철도청사",
    "(민간)광주오포D물류",
    "(민간)고려대인문관",
    "(BTL)경북대시설개선",
    "(BOT)목감휴게소신축",
    "(종심)남양주 왕숙 A-2BL",
    "(턴키)세종 5-1생활권 L5",
    "(기본)킨텍스전시장",
    "(BTL)논산관사",
    "(BTL)충남대 세종캠퍼스",
    "(BTL)공주대 세종캠퍼스",
    "(민참.분양)부천역곡대장A",
    "(민참.분양)광명시흥A",
    "(실시)부산지방합동청사",
    "(민간)KT대전인재개발원(전기)",
    "(민간)KT대전인재개발원(소방)",
    "(턴키)광교공공지산건립",
    "(시행도급)대전죽동오피스텔",
    "(시행도급)천안백석지산",
    "(시행도급)속초조양동생숙",
    "(자체.시공)인천검단1BL",
    "(자체.시공)인천검단3BL",
    "(시행도급)성남복정3BL",
    "(자체.시공)창원명곡1BL",
    "(자체.시공)시흥거모3BL",
    "(민임)이천중리4BL",
    "(시행도급)천안 오룡경기장",
    "(자체.시공)이천부필리A물류",
    "(자체.시공)서울역삼오피스텔",
    "(시행도급)성남복정홍보관",
    "(자체.시행)인천검단1BL",
    "(자체.시행)인천검단3BL",
    "(자체.시행)창원명곡1BL",
    "(자체.시행)서울역삼오피스텔",
    "(자체.시행)이천부필리A물류",
    "(자체.공모)오산운암뜰",
    "(종평)청주공공하수",
    "(종심)창녕밀양5",
    "(순수내역)포승-평택 2공구",
    "(턴키)이천문경철도8",
    "(턴키)삼성동탄철도2",
    "(턴키)안성구리10",
    "(턴키)안성구리11",
    "(종심)안성용인5",
    "(턴키)안성구리14",
    "(턴키)진접복선전철1",
    "(턴키)별내선복선전철3",
    "(적격)미금역외승강설비",
    "(적격)평택이화공공하수",
    "(적격)화천사내하수관로",
    "(적격)화천사내상수병설",
    "(종심)현대ENF천연가스",
    "(적격)서산씨지앤대산천연가스",
    "(종평)연금금성도로",
    "(종평)판교 하수처리용량 증설",
    "(적격)원산도공공하수",
    "(턴키)부산북항배후도로",
    "(실시)새만금전주8",
    "(적격)구리갈매송전관로",
    "(종평)매리양산도로",
    "(종심)포항안동2도로",
    "(종심)월곶판교전철5",
    "(종심)중리천리1도로",
    "(종심)대청댐광역상수도2",
    "(종평)울산미포산단",
    "(종평)송도기반시설",
    "(적격)장흥기산천재해예방",
    "(적격)송도배전간선",
    "(종심)새만금지구 3공구",
    "(종심)과천지식정보타운역 노반신설",
    "(종심)춘천-속초(6공구)",
    "(종심)강릉-제진(7공구)",
    "(실시)동면-진천(2공구)",
    "(턴키)용담댐 안전성 강화",
    "(실시)함양창녕3",
    "(종심)함양창녕8",
    "(턴키)삼성동탄철도4",
    "(턴키)동탄인덕원전철9",
    "(턴키)광교호매실전철2",
    "(턴키)대전철도중정비",
    "(CM)성남복정단지조성",
    "(CM)구리갈매단지조성",
    "(CM)시흥거모단지조성",
    "(CM)남양주왕숙조성2",
    "(턴키)안동댐안전성강화",
    "(기본)영동대로2도로",
    "(턴키)강북정수장증설",
    "(적격)성남복정전력구",
    "(해외)인도네시아도수관로",
    "(적격)화성능동조경",
    "(적격)파주운정조경",
    "(적격)경기 지방정원 조성",
    "(자체.시공)이천부필리도로",
    "(자체)화성우정 산단",
    "(해외)캄보디아지방도",
    "(통합)21년 창원부산",
    "(통합)23년 대구순환",
    "(통합)24년 상주영천(추가)",
    "(도로)22년 서울외곽",
    "(도로)22년 대구부산",
    "(도로)23년 신공항HW",
    "(도로)16년 광주원주",
    "(영업)23년 평택시흥",
    "(영업)22년 서울문산",
    "(ITS)25년 대전충남",
    "(ITS)25년 전북",
    "(통합)22년 상주영천",
    "(통합)24년 덕송내각",
    "(통합)24년 창원부산",
    "(도로)25년 서울문산",
    "(통합)25년 대구순환",
    "(도로)25년 서울외곽",
    "(도로)25년 대구부산",
    "(통합)21년 서부간선",
    "(통합)18년 옥산오창",
    "(도전)22년 시흥지사TN",
    "(도운)22년 인제양양TN",
    "(도운)22년 육십령TN",
    "(도전)22년 군위지사TN",
    "(국운)22년 금산TN",
    "(지운)22년 용진TN",
    "(국운)22년 밤재TN",
    "(도전)23년 동서울지사TN",
    "(지운)25년 법기TN",
    "(도전)22년 춘천지사TN",
    "(도운)22년 재약산TN",
    "(도전)23년 홍천지사TN",
    "(도전)23년 인천지사TN",
    "(국운)23년 마산TN",
    "(민운)22년 서울춘천TN",
    "(국운)23년 멧둔재TN",
    "(국운)23년 마석TN",
    "(국운)23년 김해TN",
    "(지운)23년 경기남부TN",
    "(국운)24년 백마TN",
    "(국운)23년 진주권역TN",
    "(지운)24년 중원TN",
    "(지운)25년 안민TN",
    "(국운)24년 고덕TN",
    "(도운)24년 남한산성TN",
    "(국운)25년 금산TN",
    "(지운)25년 영월저류지",
    "(지운)25년 용진TN",
    "(국운)25년 밤재TN",
    "(도운)25년 육십령TN",
    "(도전)25년 구례지사TN",
    "(도전)25년 함평지사TN",
    "(도전)25년 청송지사TN",
    "(도전)25년 군위지사TN",
    "(도전)25년 화성지사TN",
    "(도운)25년 신불산TN",
    "(도전)25년 밀양지사TN",
    "(도전)25년 창원지사TN",
    "(도전)25년 진주지사TN",
    "(민운)25년 서울춘천TN",
    "(국운)26년 멧둔재TN",
    "(BTO)화성오산고속화도로",
    "(종심)가산가평천연가스",
    "(민참.분양)하남교산A",
    "(기본)여주시신청사",
]


def load_master_names(url: str = None) -> tuple[list, str, Optional[str]]:
    """마스터 명칭 목록을 로드한다.

    Args:
        url: 서버 CSV URL. `None`이면 내장 목록 사용.

    Returns:
        `(명칭 목록, 소스 구분, DB 기준일)` 튜플.
    """
    global _LAST_MASTER_LOAD_ERROR

    _LAST_MASTER_LOAD_ERROR = ""
    fallback_names = _sanitize_master_names(DEFAULT_MASTER_NAMES)

    if not url:
        return fallback_names, "내장", None

    try:
        names, db_date = fetch_master_from_server(url)
        return names, "서버", db_date
    except Exception as e:
        _LAST_MASTER_LOAD_ERROR = str(e)
        return fallback_names, "내장", None


def get_last_master_load_error() -> str:
    """최근 마스터 로드 실패 메시지를 반환한다."""
    return _LAST_MASTER_LOAD_ERROR


# ═══════════════════════════════════════════════════════════
#  명칭 매칭 엔진 v3.2
# ═══════════════════════════════════════════════════════════
class NameMatcher:
    """공사명칭 후보를 마스터 DB와 비교해 판정한다."""

    # 공사명 후보로 볼 필요가 없는 공통 헤더/라벨 텍스트
    EXCLUDE_WORDS = {
        "공사명", "사업명", "현장명", "프로젝트명", "프로젝트",
        "일치여부", "일치", "불일치", "판정", "결과",
        "공사명칭", "명칭", "비고", "구분", "번호",
        "No", "no", "NO", "합계", "소계", "총계",
    }

    # 보고서에서 공사명 뒤에 붙는 노이즈 괄호 패턴
    NOISE_SUFFIX_RE = re.compile(
        r'\((?:LH|시공|시행|토목|건축|기계|조경|감리|설계|PM|발주|원청)\)$'
    )

    def __init__(self, master_names: list):
        """매칭에 필요한 인덱스/정규식을 초기화한다.

        Args:
            master_names: 공식명칭 목록.
        """
        self.master_names = list(dict.fromkeys(master_names))
        self.master_set = set(self.master_names)

        self.bare_to_official = build_bare_name_index(self.master_names)
        self.bare_names = list(self.bare_to_official.keys())
        self.min_bare_length = min((len(name) for name in self.bare_names), default=0)

        self.official_prefix = {
            official: split_official_name(official)[0] for official in self.master_names
        }

        self.master_suffixes = set()
        for bare in self.bare_names:
            suffix_match = re.search(r'\([^)]+\)$', bare)
            if suffix_match:
                self.master_suffixes.add(suffix_match.group())

        sorted_bares = sorted(self.bare_names, key=len, reverse=True)
        sorted_officials = sorted(self.master_names, key=len, reverse=True)
        self._bare_pattern = self._compile_alternation(sorted_bares, with_token_boundary=True)
        self._official_pattern = self._compile_alternation(sorted_officials)
        self._similarity_cache = {}
        self._bare_prefix_index: dict[str, list[str]] = {}
        for bare in self.bare_names:
            if len(bare) >= 2:
                self._bare_prefix_index.setdefault(bare[:2], []).append(bare)

    @staticmethod
    def _compile_alternation(items: list, with_token_boundary: bool = False):
        """문자열 목록을 OR 정규식으로 컴파일한다."""
        escaped_items = [re.escape(item) for item in items if item]
        if not escaped_items:
            return re.compile(r"(?!x)x")
        body = "|".join(escaped_items)
        if with_token_boundary:
            body = (
                rf'(?<![0-9A-Za-z가-힣])(?:{body})(?![0-9A-Za-z가-힣])'
            )
        return re.compile(body)

    @staticmethod
    def _remove_prefix(name: str) -> str:
        """앞쪽 접두어를 제거한 순수명칭을 반환한다."""
        _, bare = split_official_name(name)
        return bare

    @staticmethod
    def _has_prefix(name: str) -> bool:
        """문자열이 `(접두어)`로 시작하는지 확인한다."""
        return bool(PREFIX_PATTERN.match(name.strip()))

    def _normalize(self, text: str) -> str:
        """비교를 위해 텍스트를 정규화한다."""
        _, result = split_official_name(text.strip())

        suffix_match = re.search(r'\([^)]+\)$', result)
        if suffix_match:
            suffix = suffix_match.group()
            if suffix not in self.master_suffixes:
                result = result[:suffix_match.start()].strip()

        return result

    def _is_excluded(self, text: str) -> bool:
        """후보 제외 대상 텍스트인지 판정한다."""
        clean = text.strip()
        if clean in self.EXCLUDE_WORDS:
            return True
        if len(clean) <= 2:
            return True
        if clean.replace(' ', '').replace('-', '').replace('.', '').isdigit():
            return True
        return False

    def _find_containing_matches(self, normalized: str) -> list:
        """포함 관계 기반 후보 공식명칭 목록을 찾는다."""
        matches = []
        for bare, officials in self.bare_to_official.items():
            if len(bare) < 4 or len(normalized) < 4:
                continue
            if bare.startswith(normalized) or normalized.startswith(bare):
                for off in officials:
                    if off not in matches:
                        matches.append(off)
        return matches

    @staticmethod
    def _contains_bare_token(text: str, bare: str) -> bool:
        """순수명칭이 독립 토큰으로 등장하는지 검사한다."""
        token_re = re.compile(
            rf'(?<![0-9A-Za-z가-힣]){re.escape(bare)}(?![0-9A-Za-z가-힣])'
        )
        return bool(token_re.search(text))

    @staticmethod
    def _extract_prefixed_candidate(text: str, bare: str):
        """텍스트에서 `(접두어)순수명칭` 형태 후보를 추출한다."""
        prefixed_re = re.compile(
            rf'(\([^)]+\)\s*{re.escape(bare)})(?![0-9A-Za-z가-힣])'
        )
        m = prefixed_re.search(text)
        if m:
            return m.group(1).strip()
        return None

    @staticmethod
    def _is_side_boundary(text: str, index: int, is_left: bool) -> bool:
        """공식명칭 좌/우 경계가 정상인지 확인한다."""
        if is_left and index <= 0:
            return True
        if (not is_left) and index >= len(text):
            return True
        ch = text[index - 1] if is_left else text[index]
        return ch.isspace()

    def _check_official_inclusion(self, text: str) -> Optional[dict]:
        """공식명칭이 텍스트에 직접 포함되면 일치로 판정한다."""
        for match in self._official_pattern.finditer(text):
            official = match.group(0)
            return {
                "input": text,
                "status": "일치",
                "suggestion": official,
                "issue": "",
            }
        return None

    @staticmethod
    def _format_ambiguous_issue(candidates: list[str]) -> str:
        """특정불가 사유 문자열을 포맷팅한다."""
        return f"공사명 불일치: 특정불가 (유사 {len(candidates)}건)"

    @staticmethod
    def _build_ambiguous_result(text: str, candidates: list[str]) -> dict:
        """특정불가 후보들을 하나의 결과 행으로 변환한다."""
        return {
            "input": text,
            "status": "불일치",
            "suggestion": " / ".join(candidates),
            "issue": NameMatcher._format_ambiguous_issue(candidates),
        }

    def find_all_in_text(self, text: str) -> list:
        """텍스트에서 마스터DB와 연관된 후보 문자열을 찾는다."""
        found = []
        text = text.strip()
        if len(text) < 3:
            return []
        if self._is_excluded(text):
            return []

        if text in self.master_set:
            return [text]
        if text in self.bare_to_official:
            return [text]

        norm = self._normalize(text)
        if norm != text and norm in self.bare_to_official:
            return [text]
        if len(text) < self.min_bare_length:
            return []

        for match in self._official_pattern.finditer(text):
            official = match.group(0)
            if official not in found:
                found.append(official)

        for match in self._bare_pattern.finditer(text):
            bare = match.group(0)
            officials = self.bare_to_official.get(bare)
            if not officials:
                continue

            already = any(off in found for off in officials)
            if already:
                continue

            prefixed = self._extract_prefixed_candidate(text, bare)
            if prefixed:
                if prefixed not in found:
                    found.append(prefixed)
                continue

            if bare not in found:
                found.append(bare)

        if not found:
            containing = self._find_containing_matches(norm)
            if containing:
                return [text]

            best_name, best_score = self._best_similarity(norm)
            if best_score >= 0.7:
                return [text]

        return found

    def _best_similarity(self, text: str):
        """순수명칭 기준 최고 유사도 공식명칭을 찾는다."""
        cached = self._similarity_cache.get(text)
        if cached is not None:
            return cached

        best_name = None
        best_score = 0.0

        candidates = []
        if len(text) >= 2:
            candidates = self._bare_prefix_index.get(text[:2], [])
        if not candidates and len(text) >= 1:
            for key, bares in self._bare_prefix_index.items():
                if key and key[0] == text[0]:
                    candidates.extend(bares)

        if candidates:
            for bare in candidates:
                score = SequenceMatcher(None, text, bare).ratio()
                if score > best_score:
                    best_score = score
                    best_name = self.bare_to_official.get(bare, [None])[0]

        if best_score < 0.7:
            for bare, officials in self.bare_to_official.items():
                score = SequenceMatcher(None, text, bare).ratio()
                if score > best_score:
                    best_score = score
                    best_name = officials[0]

        result = (best_name, best_score)
        self._similarity_cache[text] = result
        return result

    def match(self, text: str) -> Optional[dict]:
        """단일 문자열을 마스터 명칭 규칙으로 판정한다.

        판정 결과의 불일치 사유는 3가지 카테고리로 분류된다:
        - 공사명 불일치: 순수명칭 자체가 다름
        - 접두어(사업유형) 불일치: 순수명칭은 맞으나 접두어가 다름
        - 연도 불일치: 접두어·공사명은 같고 연도만 다름

        Args:
            text: 검사 대상 문자열.

        Returns:
            ``{input, status, suggestion, issue}`` 형태 dict 또는 ``None``.
        """
        text = str(text or "").strip()
        if not text or self._is_excluded(text):
            return None

        match_text = _normalize_known_prefix_punctuation(text)

        # ── STEP 1: 공식명칭 포함 판정 (일치) ──
        inclusion_result = self._check_official_inclusion(match_text)
        if inclusion_result:
            inclusion_result["input"] = text
            return inclusion_result

        # ── STEP 2: 접두어 분리 매칭 ──
        normalized = self._normalize(match_text)
        split_prefix, split_bare = strip_known_prefix(match_text)
        split_bare = split_bare or normalized
        stripped_input = match_text.strip().strip(STRIP_CHARS).strip()
        is_direct_prefix_form = (
            bool(split_prefix)
            and not self._has_prefix(match_text)
            and stripped_input.startswith(split_prefix)
        )

        if split_bare in self.bare_to_official:
            officials = self.bare_to_official[split_bare]
            if len(officials) == 1:
                official = officials[0]
                official_prefix = self.official_prefix.get(official, "")

                if not split_prefix:
                    return {
                        "input": text,
                        "status": "불일치",
                        "suggestion": official,
                        "issue": f"접두어(사업유형) 불일치: 접두어 누락 → 공식: {official}",
                    }

                if split_prefix == official_prefix:
                    return {
                        "input": text,
                        "status": "불일치",
                        "suggestion": official,
                        "issue": f"접두어(사업유형) 불일치: 괄호 표기 누락 → 공식: {official}",
                    }

                return {
                    "input": text,
                    "status": "불일치",
                    "suggestion": official,
                    "issue": (
                        f"접두어(사업유형) 불일치: "
                        f"{split_prefix} → {official_prefix} → 공식: {official}"
                    ),
                }

            return self._build_ambiguous_result(text, officials)

        # 괄호 없는 "접두어+명칭" 형태의 유사 매칭
        if is_direct_prefix_form and len(split_bare) >= 4:
            best_split_name, best_split_score = self._best_similarity(split_bare)
            if (
                best_split_name
                and best_split_score >= 0.9
                and self.official_prefix.get(best_split_name, "") == split_prefix
            ):
                return {
                    "input": text,
                    "status": "불일치",
                    "suggestion": best_split_name,
                    "issue": (
                        f"접두어(사업유형) 불일치: "
                        f"괄호 표기 누락 → 공식: {best_split_name}"
                    ),
                }

        # ── STEP 3: 포함 관계 비교 ──
        containing = self._find_containing_matches(normalized)
        if containing:
            if len(containing) == 1:
                _, candidate_bare = split_official_name(containing[0])
                candidate_bare = candidate_bare.replace(" ", "")
                compare_text = normalized.replace(" ", "")

                likely_typo = (
                    compare_text.startswith(candidate_bare)
                    and len(compare_text) - len(candidate_bare) <= 1
                )
                if not likely_typo:
                    return {
                        "input": text,
                        "status": "불일치",
                        "suggestion": containing[0],
                        "issue": f"공사명 불일치: 명칭 불완전 → 공식: {containing[0]}",
                    }
            else:
                return self._build_ambiguous_result(text, containing)

        # ── STEP 4: 유사도 비교 ──
        best_name, best_score = self._best_similarity(normalized)
        if best_score >= 0.7:
            _, best_bare = split_official_name(best_name)
            best_prefix = self.official_prefix.get(best_name, "")

            # 연도만 다른지 판별 (접두어 분리된 경우)
            if (
                split_prefix
                and split_prefix == best_prefix
                and _is_year_only_diff(split_bare, best_bare)
            ):
                return {
                    "input": text,
                    "status": "불일치",
                    "suggestion": best_name,
                    "issue": f"연도 불일치 → 유사: {best_name}",
                }

            # 연도만 다른지 판별 (정규화된 전체 비교)
            if _is_year_only_diff(normalized, best_bare):
                return {
                    "input": text,
                    "status": "불일치",
                    "suggestion": best_name,
                    "issue": f"연도 불일치 → 유사: {best_name}",
                }

            pct = f"{best_score * 100:.0f}%"
            return {
                "input": text,
                "status": "불일치",
                "suggestion": best_name,
                "issue": f"공사명 불일치 (유사도 {pct}) → 유사: {best_name}",
            }

        return None

    def check(self, text: str):
        """단일 문자열을 마스터 명칭 규칙으로 판정한다."""
        return self.match(text)

# ═══════════════════════════════════════════════════════════
#  검토 엔진
# ═══════════════════════════════════════════════════════════
class ReviewEngine:
    """파일 단위로 후보 명칭을 수집하고 판정 결과를 집계한다."""

    def __init__(self, matcher: NameMatcher):
        """검토 엔진을 초기화한다."""
        self.matcher = matcher

    @staticmethod
    def _build_full_text_with_offsets(text_items: list) -> tuple[str, list, list]:
        """텍스트 조각 목록을 단일 문자열과 오프셋 테이블로 변환한다."""
        parts = []
        starts = []
        offsets = []
        cursor = 0

        for location, raw_text in text_items:
            text = str(raw_text).strip()
            if not text:
                continue

            parts.append(text)
            start = cursor
            end = start + len(text)
            starts.append(start)
            offsets.append((start, end, location, text))
            cursor = end + 1

        return "\n".join(parts), starts, offsets

    @staticmethod
    def _find_offset_index(starts: list, offsets: list, position: int) -> int:
        """결합 문자열 위치를 원본 조각 인덱스로 역매핑한다."""
        idx = bisect.bisect_right(starts, position) - 1
        if idx < 0 or idx >= len(offsets):
            return -1
        start, end, _, _ = offsets[idx]
        if not (start <= position < end):
            return -1
        return idx

    @staticmethod
    def _extend_official_candidate(source_text: str, start: int, end: int) -> str:
        """공식명칭 주변 부착 문자를 포함한 후보를 생성한다."""
        left = start
        right = end

        if left > 0 and not source_text[left - 1].isspace():
            left -= 1
        if right < len(source_text) and not source_text[right].isspace():
            right += 1
        return source_text[left:right].strip()

    @staticmethod
    def _iter_aux_candidates(text: str) -> list[str]:
        """보조 탐색용 후보 문자열을 생성한다."""
        primary = _strip_surrounding(text)
        return [primary] if primary else []

    def review_file(self, filepath: str, progress_callback=None) -> dict:
        """단일 파일을 검토해 공사명칭 판정 결과를 반환한다."""
        def _emit_progress(value: float):
            if progress_callback is None:
                return
            try:
                clamped = max(0.0, min(1.0, float(value)))
                progress_callback(clamped)
            except Exception:
                pass

        filename = os.path.basename(filepath)
        try:
            text_items = extract_text_from_file(
                filepath,
                progress_callback=progress_callback,
            )
            _emit_progress(0.3)
        except Exception as e:
            _emit_progress(1.0)
            return {
                "file": filename, "path": filepath,
                "total": 0, "matched": 0, "mismatched": 0, "warning": 0,
                "overall": "오류", "details": [], "error": str(e)
            }

        results = []
        checked_results: dict[tuple[str, str, str, str], dict] = {}
        checked_locations: dict[tuple[str, str, str, str], list[str]] = {}

        def _location_scope_key(location: str) -> str:
            """결과 병합용 위치 범위를 계산한다."""
            page_match = re.search(r'P(\d+)', location or "")
            if page_match:
                return f"P{page_match.group(1)}"
            return str(location or "")

        def consume_candidate(candidate: str, location: str):
            match_result = self.matcher.match(candidate)
            if not match_result:
                return

            if isinstance(match_result, list):
                match_results = match_result
            else:
                match_results = [match_result]

            for match_item in match_results:
                key = (
                    _location_scope_key(location),
                    match_item.get("input", candidate),
                    match_item.get("status", ""),
                    match_item.get("suggestion", ""),
                )

                if key in checked_results:
                    existing = checked_results[key]
                    locs = checked_locations[key]
                    if location not in locs:
                        locs.append(location)
                        existing["location"] = ", ".join(locs)
                    continue

                stored_result = dict(match_item)
                stored_result["location"] = location
                results.append(stored_result)
                checked_results[key] = stored_result
                checked_locations[key] = [location]

        full_text, starts, offsets = self._build_full_text_with_offsets(text_items)
        if full_text:
            match_events = []
            matched_offset_indices = set()
            aux_candidate_cache: dict[str, list[str]] = {}
            prefix_cache: dict[str, Optional[str]] = {}
            has_korean_re = re.compile(r"[가-힣]")
            min_aux_length = self.matcher.min_bare_length * 0.5
            page_lines = full_text.splitlines()

            for match in self.matcher._official_pattern.finditer(full_text):
                idx = self._find_offset_index(starts, offsets, match.start())
                if idx < 0:
                    continue
                matched_offset_indices.add(idx)
                line_start, _, location, source_text = offsets[idx]
                local_start = match.start() - line_start
                local_end = local_start + len(match.group(0))
                candidate = self._extend_official_candidate(source_text, local_start, local_end)
                match_events.append((match.start(), 0, candidate, location))

            for match in self.matcher._bare_pattern.finditer(full_text):
                idx = self._find_offset_index(starts, offsets, match.start())
                if idx < 0:
                    continue
                matched_offset_indices.add(idx)
                _, _, location, source_text = offsets[idx]
                bare = match.group(0)
                prefixed = self.matcher._extract_prefixed_candidate(source_text, bare)
                if prefixed:
                    candidate = prefixed
                else:
                    bare_line_idx = full_text[:match.start()].count("\n")
                    cache_key = f"{bare}|{bare_line_idx}"
                    if cache_key in prefix_cache:
                        reconstructed = prefix_cache[cache_key]
                    else:
                        reconstructed = _reconstruct_prefix(
                            bare, page_lines, bare_line_idx,
                            self.matcher.master_set, KNOWN_PREFIXES
                        )
                        prefix_cache[cache_key] = reconstructed
                    candidate = reconstructed if reconstructed else bare
                match_events.append((match.start(), 1, candidate, location))

            _emit_progress(0.6)
            match_events.sort(key=lambda item: (item[0], item[1]))

            for _, _, candidate, location in match_events:
                consume_candidate(candidate, location)

            unmatched_indices = [
                idx for idx in range(len(offsets))
                if idx not in matched_offset_indices
            ]
            unmatched_total = len(unmatched_indices)

            for processed, idx in enumerate(unmatched_indices, 1):
                _, _, location, source_text = offsets[idx]
                if processed % 100 == 0 or processed == unmatched_total:
                    ratio = processed / unmatched_total if unmatched_total else 1.0
                    _emit_progress(0.6 + 0.35 * ratio)

                if idx in matched_offset_indices:
                    continue
                text = source_text.strip()
                if len(text) > 80:
                    continue
                if len(text) < 2:
                    continue
                if text in self.matcher.EXCLUDE_WORDS:
                    continue
                if not has_korean_re.search(text):
                    continue
                if min_aux_length > 0 and len(text) < min_aux_length:
                    continue

                consumed_probe = False
                for probe_text in self._iter_aux_candidates(text):
                    if min_aux_length > 0 and len(probe_text) < min_aux_length:
                        continue
                    if probe_text in self.matcher.EXCLUDE_WORDS:
                        continue
                    if not has_korean_re.search(probe_text):
                        continue

                    cached_candidates = aux_candidate_cache.get(probe_text)
                    if cached_candidates is None:
                        cached_candidates = self.matcher.find_all_in_text(probe_text)
                        aux_candidate_cache[probe_text] = cached_candidates

                    for candidate in cached_candidates:
                        consume_candidate(candidate, location)
                        consumed_probe = True

                if consumed_probe or "(" in text or ")" in text:
                    continue

                parts = text.split()
                if not parts:
                    continue
                lead = _strip_surrounding(parts[0])
                if min_aux_length > 0 and len(lead) < min_aux_length:
                    continue
                if lead in self.matcher.EXCLUDE_WORDS:
                    continue
                if not has_korean_re.search(lead):
                    continue

                lead_result = self.matcher.match(lead)
                if (
                    lead_result
                    and "특정불가" in str(lead_result.get("issue", ""))
                ):
                    consume_candidate(lead, location)
        else:
            _emit_progress(0.6)

        matched = sum(1 for r in results if r["status"] == "일치")
        mismatched = sum(1 for r in results if r["status"] == "불일치")
        warnings = sum(
            1 for r in results if r["status"] not in {"일치", "불일치"}
        )
        total = len(results)

        if total == 0:
            overall = "명칭없음"
        elif mismatched == 0 and warnings == 0:
            overall = "적합"
        else:
            overall = "검토필요"

        def _sort_key(result: dict) -> tuple[int, int]:
            loc = result.get("location", "")
            match = re.match(r'P(\d+)\s*L(\d+)', loc)
            if match:
                return int(match.group(1)), int(match.group(2))
            return (999999, 999999)

        results.sort(key=_sort_key)
        _emit_progress(1.0)
        return {
            "file": filename, "path": filepath,
            "total": total, "matched": matched,
            "mismatched": mismatched, "warning": warnings, "overall": overall,
            "details": results, "error": None
        }


# ═══════════════════════════════════════════════════════════
#  Excel 리포트 생성
# ═══════════════════════════════════════════════════════════
def save_excel_report(all_results: list, output_path: str,
                      master_names: list):
    """검토 결과를 Excel 리포트로 저장한다."""
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

    wb = Workbook()
    hdr_fill = PatternFill("solid", fgColor="2F5496")
    hdr_font = Font(color="FFFFFF", bold=True, size=11, name="맑은 고딕")
    ng_fill = PatternFill("solid", fgColor="FFC7CE")
    warn_fill = PatternFill("solid", fgColor="FFF2CC")
    bdr = Border(left=Side('thin'), right=Side('thin'),
                 top=Side('thin'), bottom=Side('thin'))
    bfont = Font(size=10, name="맑은 고딕")
    ctr = Alignment(horizontal='center', vertical='center')

    ws = wb.active
    ws.title = "불일치목록"

    ws.merge_cells('A1:E1')
    ws['A1'].value = (
        f"공사명칭 불일치 검토 리포트  |  "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    ws['A1'].font = Font(bold=True, size=14, name="맑은 고딕")
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 35

    t_found = sum(r["total"] for r in all_results)
    t_match = sum(r["matched"] for r in all_results)
    t_mis = sum(r["mismatched"] for r in all_results)

    ws.merge_cells('A2:E2')
    ws['A2'].value = (
        f"파일 {len(all_results)}개  |  "
        f"발견 {t_found}개  |  "
        f"일치 {t_match}개  |  "
        f"불일치 {t_mis}개"
    )
    ws['A2'].font = Font(size=11, name="맑은 고딕", bold=True)
    ws['A2'].alignment = Alignment(horizontal='center')

    headers = ["파일명", "위치", "보고서 기재 명칭",
               "공식명칭(추천)", "불일치 사유"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = ctr
        cell.border = bdr

    row = 5
    for fr in all_results:
        if fr.get("error"):
            ws.cell(row=row, column=1, value=fr["file"]).font = bfont
            ws.cell(row=row, column=3, value="오류").font = bfont
            ws.cell(row=row, column=5, value=fr["error"]).font = bfont
            for c in range(1, 6):
                ws.cell(row=row, column=c).border = bdr
            row += 1
            continue

        ng_items = [
            d for d in fr["details"]
            if d["status"] in ("불일치", "경고")
        ]

        if not ng_items and not fr.get("error"):
            continue

        for d in ng_items:
            ws.cell(row=row, column=1, value=fr["file"]).font = bfont
            ws.cell(row=row, column=2,
                    value=d.get("location", "")).font = bfont
            c3 = ws.cell(row=row, column=3, value=d["input"])
            c3.font = bfont

            if d["status"] == "경고":
                c3.fill = warn_fill
            else:
                c3.fill = ng_fill

            ws.cell(row=row, column=4,
                    value=d.get("suggestion", "")).font = bfont
            ws.cell(row=row, column=5,
                    value=d.get("issue", "")).font = bfont
            for c in range(1, 6):
                ws.cell(row=row, column=c).border = bdr
            row += 1

    for letter, w in zip('ABCDE', [30, 18, 35, 40, 55]):
        ws.column_dimensions[letter].width = w

    ws2 = wb.create_sheet("마스터목록")
    for ci, h in enumerate(["No.", "공식 명칭", "접두어", "순수 명칭"], 1):
        cell = ws2.cell(row=1, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = ctr
        cell.border = bdr

    for i, name in enumerate(master_names, 1):
        prefix, bare = split_official_name(name)
        ws2.cell(row=i + 1, column=1, value=i).font = bfont
        ws2.cell(row=i + 1, column=2, value=name).font = bfont
        ws2.cell(row=i + 1, column=3, value=prefix).font = bfont
        ws2.cell(row=i + 1, column=4, value=bare).font = bfont
        for c in range(1, 5):
            ws2.cell(row=i + 1, column=c).border = bdr

    for letter, w in zip('ABCD', [8, 45, 15, 35]):
        ws2.column_dimensions[letter].width = w

    wb.save(output_path)


def generate_highlight_snapshots(
    filepath: str,
    ng_items: list[dict],
    resolution: int = 250,
) -> list[tuple[int, bytes]]:
    """PDF 파일에서 불일치 위치를 하이라이트한 페이지 이미지를 생성한다.

    PyMuPDF로 순수명칭(접두어·연도 제거)을 페이지에서 직접 검색하여
    좌표를 확보하고, 해당 위치에 하이라이트 박스를 그린다.
    전체 키워드 검색 실패 시 한글만 추출하여 재검색한다.
    한글 검색 시 접두어·영문·숫자 영역을 포함하도록 좌우 패딩을 확장한다.
    """
    try:
        import io
        import fitz
    except ImportError:
        return []

    if os.path.splitext(filepath)[1].lower() != ".pdf":
        return []

    page_targets: dict[int, list[dict]] = {}
    for item in ng_items:
        location = item.get("location", "")
        if not location:
            continue
        for page_str in re.findall(r'P(\d+)', location):
            page_num = int(page_str)
            page_targets.setdefault(page_num, [])
            if item not in page_targets[page_num]:
                page_targets[page_num].append(item)

    if not page_targets:
        return []

    _year_prefix_re = re.compile(r'^\d{2,4}\s*년\s*')

    def _get_search_keyword(item: dict) -> str:
        raw_input = item.get("input", "") or ""
        raw_input = raw_input.strip().strip(STRIP_CHARS).strip()
        _, bare = strip_known_prefix(raw_input)
        if not bare:
            bare = raw_input
        bare = _year_prefix_re.sub("", bare).strip()
        return bare

    dpi_scale = resolution / 72.0
    mat = fitz.Matrix(dpi_scale, dpi_scale)

    results: list[tuple[int, bytes]] = []
    try:
        doc = fitz.open(filepath)
        try:
            for page_num in sorted(page_targets.keys()):
                page_idx = page_num - 1
                if page_idx >= len(doc):
                    continue

                page = doc[page_idx]
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")

                from PIL import Image, ImageDraw
                pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                draw = ImageDraw.Draw(pil_img)
                drew_any = False

                for item in page_targets[page_num]:
                    keyword = _get_search_keyword(item)
                    if not keyword or len(keyword) < 2:
                        continue

                    used_hangul_fallback = False
                    rects = page.search_for(keyword)

                    if not rects:
                        hangul_only = re.sub(r'[^가-힣]', '', keyword)
                        if hangul_only and len(hangul_only) >= 2:
                            rects = page.search_for(hangul_only)
                            if rects:
                                used_hangul_fallback = True

                    for rect in rects:
                        x0 = rect.x0 * dpi_scale
                        y0 = rect.y0 * dpi_scale
                        x1 = rect.x1 * dpi_scale
                        y1 = rect.y1 * dpi_scale
                        char_h = y1 - y0
                        pad_y = char_h * 0.15

                        if used_hangul_fallback:
                            pad_left = char_h * 2.5
                            pad_right = char_h * 1.5
                        else:
                            pad_left = char_h * 0.5
                            pad_right = char_h * 0.5

                        box = (x0 - pad_left, y0 - pad_y,
                               x1 + pad_right, y1 + pad_y)
                        draw.rectangle(box, outline="red", width=3)
                        drew_any = True

                if drew_any:
                    buf = io.BytesIO()
                    pil_img.save(buf, format="PNG")
                    results.append((page_num, buf.getvalue()))

        finally:
            doc.close()
    except Exception:
        pass

    return results
